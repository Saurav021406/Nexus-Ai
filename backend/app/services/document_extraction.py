"""Document text extraction (Step 2 of the RAG/Domain Router design).

    PDF/Word file -> text extract karo -> [Step 3: chunks mein todo] -> ...

This module is deliberately just extraction - no chunking, no embeddings,
no storage decisions. Those are Steps 3-4, separate pieces of work on
purpose (the RAG build order explicitly lists "Document upload + text
extraction" as its own step before "Chunking + embeddings + pgvector
setup"). Keeping this narrow means it's easy to verify in isolation: given
file bytes, does the right text come out.

Two extractors, both pure-Python, no system-level dependencies:
  - PDF: pypdf (new dependency, lightweight, no external binaries)
  - Word (.docx only, not legacy .doc): python-docx (already a dependency,
    added earlier for the Report Agent's DOCX export - reused here for
    reading instead of writing)
"""

from __future__ import annotations

import io

from docx import Document
from pypdf import PdfReader


class ExtractionError(Exception):
    """Raised when a file claims to be a PDF/DOCX but can't actually be
    parsed as one (corrupt file, password-protected PDF, wrong extension,
    etc.) - callers should turn this into a clear 400, not a 500 crash."""


def extract_pdf_text(raw_bytes: bytes) -> dict:
    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not open PDF: {e}")

    if reader.is_encrypted:
        raise ExtractionError("This PDF is password-protected and can't be read.")

    pages_text = []
    for page in reader.pages:
        try:
            pages_text.append(page.extract_text() or "")
        except Exception:
            pages_text.append("")  # one bad page shouldn't fail the whole document

    full_text = "\n\n".join(pages_text).strip()
    if not full_text:
        raise ExtractionError(
            "No extractable text found - this PDF may be scanned images without a text "
            "layer (OCR isn't supported yet)."
        )

    return {
        "text": full_text,
        "page_count": len(reader.pages),
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
    }


def extract_docx_text(raw_bytes: bytes) -> dict:
    try:
        document = Document(io.BytesIO(raw_bytes))
    except Exception as e:
        raise ExtractionError(f"Could not open Word document: {e}")

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    # Tables carry real content too (e.g. contract terms in a table) - pull
    # cell text out rather than silently dropping it.
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    full_text = "\n".join(paragraphs).strip()
    if not full_text:
        raise ExtractionError("No extractable text found in this Word document.")

    return {
        "text": full_text,
        "paragraph_count": len(paragraphs),
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
    }


def extract_document_text(raw_bytes: bytes, filename: str) -> dict:
    """Dispatches by extension. Returns the extractor's dict plus a
    "document_type" field so callers don't need to re-check the extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        result = extract_pdf_text(raw_bytes)
        result["document_type"] = "pdf"
        return result
    if lower.endswith(".docx"):
        result = extract_docx_text(raw_bytes)
        result["document_type"] = "docx"
        return result
    raise ExtractionError(f"Unsupported document type for extraction: {filename}")
