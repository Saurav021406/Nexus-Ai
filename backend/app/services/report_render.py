"""Renders Report Agent content (plain JSON) into downloadable PDF and DOCX
files. No LLM calls happen here - purely deterministic formatting, so the
same generated content can be exported to both formats without re-running
the agent or risking inconsistent numbers between the two files.
"""

from __future__ import annotations

import base64
import io
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    ListFlowable,
    ListItem,
    Table,
    TableStyle,
    Image,
    PageBreak,
)
from reportlab.graphics.shapes import Drawing, Circle, String
from reportlab.lib.enums import TA_CENTER


BRAND_COLOR = "0891b2"  # cyan-600, matches the frontend accent
BRAND_DARK = "0f172a"  # slate-900, matches the frontend background


def _cover_page(report: dict[str, Any], styles) -> list:
    """A branded title page: monogram badge, title, subtitle, meta."""
    cover: list = [Spacer(1, 1.6 * inch)]

    badge = Drawing(120, 120)
    badge.add(Circle(60, 60, 56, fillColor=colors.HexColor(f"#{BRAND_COLOR}"), strokeColor=None))
    badge.add(
        String(
            60,
            42,
            "N",
            fontName="Helvetica-Bold",
            fontSize=48,
            fillColor=colors.white,
            textAnchor="middle",
        )
    )
    badge.hAlign = "CENTER"
    cover.append(badge)
    cover.append(Spacer(1, 0.4 * inch))

    cover_title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontSize=26,
        alignment=TA_CENTER,
        textColor=colors.HexColor(f"#{BRAND_COLOR}"),
    )
    cover_subtitle_style = ParagraphStyle(
        "CoverSubtitle",
        parent=styles["Normal"],
        fontSize=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1e293b"),
        spaceBefore=8,
    )
    cover_meta_style = ParagraphStyle(
        "CoverMeta",
        parent=styles["Normal"],
        fontSize=9,
        alignment=TA_CENTER,
        textColor=colors.grey,
        spaceBefore=6,
    )

    cover.append(Paragraph(report.get("title", "Nexus AI Analysis Report"), cover_title_style))
    cover.append(Paragraph(report.get("subtitle", report.get("filename", "")), cover_subtitle_style))
    cover.append(Spacer(1, 0.3 * inch))
    cover.append(
        Paragraph(
            f"Domain: {report.get('primary_domain', 'General')}",
            cover_meta_style,
        )
    )
    cover.append(Paragraph(f"Generated: {report.get('generated_at', '')}", cover_meta_style))
    if report.get("participating_agents"):
        cover.append(
            Paragraph(
                f"Agents involved: {', '.join(report['participating_agents'])}",
                cover_meta_style,
            )
        )
    cover.append(Spacer(1, 1.5 * inch))
    cover.append(
        Paragraph("An AI-generated, multi-agent analysis report", cover_meta_style)
    )
    cover.append(PageBreak())
    return cover


def render_pdf(report: dict[str, Any]) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle", parent=styles["Title"], textColor=colors.HexColor(f"#{BRAND_COLOR}")
    )
    heading_style = ParagraphStyle(
        "ReportHeading",
        parent=styles["Heading2"],
        textColor=colors.HexColor(f"#{BRAND_COLOR}"),
        spaceBefore=14,
        spaceAfter=6,
    )
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], textColor=colors.grey, fontSize=9)
    body_style = styles["BodyText"]
    caption_style = ParagraphStyle(
        "Caption", parent=styles["Normal"], fontSize=8, textColor=colors.grey, alignment=TA_CENTER
    )

    story: list = []
    story.extend(_cover_page(report, styles))

    story.append(Paragraph("Executive Summary", heading_style))
    story.append(Paragraph(report.get("executive_summary", ""), body_style))

    if report.get("agent_collaboration_narrative"):
        story.append(Paragraph("How the Agents Reached This Conclusion", heading_style))
        story.append(Paragraph(report["agent_collaboration_narrative"], body_style))
        story.append(Spacer(1, 6))

        timeline = report.get("agent_timeline", [])
        if timeline:
            cell_style = ParagraphStyle("TimelineCell", parent=styles["BodyText"], fontSize=8, leading=10)
            header_cell_style = ParagraphStyle(
                "TimelineHeader", parent=styles["BodyText"], fontSize=8, leading=10, textColor=colors.white
            )
            timeline_data = [
                [Paragraph(h, header_cell_style) for h in ("Agent", "Role", "Outcome")]
            ] + [
                [
                    Paragraph(step.get("agent", ""), cell_style),
                    Paragraph(step.get("role", ""), cell_style),
                    Paragraph(step.get("detail", ""), cell_style),
                ]
                for step in timeline
            ]
            timeline_table = Table(timeline_data, colWidths=[1.3 * inch, 1.3 * inch, 3.4 * inch])
            timeline_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BRAND_DARK}")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTSIZE", (0, 0), (-1, -1), 8),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ]
                )
            )
            story.append(timeline_table)

    if report.get("key_metrics"):
        story.append(Paragraph("Key Metrics", heading_style))
        metric_cell_style = ParagraphStyle("MetricCell", parent=styles["BodyText"], fontSize=9, leading=11)
        table_data = [["#", "Metric"]] + [
            [str(i + 1), Paragraph(metric, metric_cell_style)]
            for i, metric in enumerate(report["key_metrics"])
        ]
        table = Table(table_data, colWidths=[0.4 * inch, 5.6 * inch])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BRAND_COLOR}")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ]
            )
        )
        story.append(table)

    if report.get("recommendation"):
        story.append(Paragraph("Recommendation", heading_style))
        story.append(Paragraph(report["recommendation"], body_style))

    for section in report.get("sections", []):
        story.append(Paragraph(section["heading"], heading_style))
        if section.get("body"):
            story.append(Paragraph(section["body"], body_style))
        if section.get("bullets"):
            story.append(
                ListFlowable(
                    [ListItem(Paragraph(b, body_style)) for b in section["bullets"]],
                    bulletType="bullet",
                )
            )
        if section.get("recommendation"):
            story.append(Spacer(1, 4))
            story.append(Paragraph(f"<b>Recommendation:</b> {section['recommendation']}", body_style))

    charts = report.get("charts", [])
    if charts:
        story.append(Paragraph("Visualizations", heading_style))
        for chart in charts:
            try:
                image_bytes = base64.b64decode(chart["image_base64"])
                img = Image(io.BytesIO(image_bytes), width=5.5 * inch, height=3.7 * inch)
                img.hAlign = "CENTER"
                story.append(img)
                story.append(Paragraph(chart.get("title", ""), caption_style))
                story.append(Spacer(1, 10))
            except Exception as e:
                print(f"Skipping chart in PDF export: {e}")

    review = report.get("review")
    security = report.get("security")
    if review or security:
        story.append(Paragraph("Quality & Safety Review", heading_style))
        if review:
            story.append(
                Paragraph(
                    f"Overall quality: <b>{review.get('overall_quality', 'n/a')}</b> · "
                    f"Approved: <b>{review.get('approved', 'n/a')}</b>",
                    body_style,
                )
            )
        if security:
            story.append(
                Paragraph(
                    f"Risk level: <b>{security.get('risk_level', 'n/a')}</b> · "
                    f"Safe to show: <b>{security.get('safe_to_show', 'n/a')}</b>",
                    body_style,
                )
            )

    story.append(Spacer(1, 20))
    story.append(Paragraph("Generated by Nexus AI", meta_style))

    doc.build(story)
    return buffer.getvalue()


def render_docx(report: dict[str, Any]) -> bytes:
    document = Document()

    # --- Cover page ---
    cover_title = document.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        document.add_paragraph()
    badge_p = document.add_paragraph()
    badge_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    badge_run = badge_p.add_run("●")
    badge_run.font.size = Pt(54)
    badge_run.font.color.rgb = RGBColor.from_string(BRAND_COLOR)

    cover_title_p = document.add_paragraph()
    cover_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title_run = cover_title_p.add_run(report.get("title", "Nexus AI Analysis Report"))
    cover_title_run.font.size = Pt(26)
    cover_title_run.bold = True
    cover_title_run.font.color.rgb = RGBColor.from_string(BRAND_COLOR)

    cover_subtitle_p = document.add_paragraph()
    cover_subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_subtitle_run = cover_subtitle_p.add_run(report.get("subtitle", report.get("filename", "")))
    cover_subtitle_run.font.size = Pt(14)

    cover_meta_p = document.add_paragraph()
    cover_meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_meta_run = cover_meta_p.add_run(
        f"Domain: {report.get('primary_domain', 'General')}\n"
        f"Generated: {report.get('generated_at', '')}"
        + (
            f"\nAgents involved: {', '.join(report['participating_agents'])}"
            if report.get("participating_agents")
            else ""
        )
    )
    cover_meta_run.font.size = Pt(9)
    cover_meta_run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

    document.add_page_break()

    # --- Body ---
    title = document.add_heading(report["title"], level=0)
    title.runs[0].font.color.rgb = RGBColor.from_string(BRAND_COLOR)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(
        f"Domain: {report.get('primary_domain', 'General')}    "
        f"Generated: {report.get('generated_at', '')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x64, 0x64, 0x64)

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(report.get("executive_summary", ""))

    if report.get("agent_collaboration_narrative"):
        document.add_heading("How the Agents Reached This Conclusion", level=1)
        document.add_paragraph(report["agent_collaboration_narrative"])

        timeline = report.get("agent_timeline", [])
        if timeline:
            table = document.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = "Agent", "Role", "Outcome"
            for step in timeline:
                row = table.add_row().cells
                row[0].text = step.get("agent", "")
                row[1].text = step.get("role", "")
                row[2].text = step.get("detail", "")

    if report.get("key_metrics"):
        document.add_heading("Key Metrics", level=1)
        table = document.add_table(rows=1, cols=1)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Metric"
        for metric in report["key_metrics"]:
            row = table.add_row()
            row.cells[0].text = metric

    if report.get("recommendation"):
        document.add_heading("Recommendation", level=1)
        document.add_paragraph(report["recommendation"])

    for section in report.get("sections", []):
        document.add_heading(section["heading"], level=1)
        if section.get("body"):
            document.add_paragraph(section["body"])
        for bullet in section.get("bullets", []):
            document.add_paragraph(bullet, style="List Bullet")
        if section.get("recommendation"):
            p = document.add_paragraph()
            p.add_run("Recommendation: ").bold = True
            p.add_run(section["recommendation"])

    charts = report.get("charts", [])
    if charts:
        document.add_heading("Visualizations", level=1)
        for chart in charts:
            try:
                image_bytes = base64.b64decode(chart["image_base64"])
                document.add_picture(io.BytesIO(image_bytes), width=Inches(5.5))
                caption = document.add_paragraph(chart.get("title", ""))
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].font.size = Pt(8)
                caption.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            except Exception as e:
                print(f"Skipping chart in DOCX export: {e}")

    review = report.get("review")
    security = report.get("security")
    if review or security:
        document.add_heading("Quality & Safety Review", level=1)
        if review:
            document.add_paragraph(
                f"Overall quality: {review.get('overall_quality', 'n/a')}  ·  "
                f"Approved: {review.get('approved', 'n/a')}"
            )
        if security:
            document.add_paragraph(
                f"Risk level: {security.get('risk_level', 'n/a')}  ·  "
                f"Safe to show: {security.get('safe_to_show', 'n/a')}"
            )

    footer = document.add_paragraph()
    footer_run = footer.add_run("Generated by Nexus AI")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
