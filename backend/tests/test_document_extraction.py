import io

import pytest
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

from app.services.document_extraction import (
    ExtractionError,
    extract_document_text,
    extract_docx_text,
    extract_pdf_text,
)


def _make_test_pdf(lines_per_page: list[list[str]]) -> bytes:
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    for lines in lines_per_page:
        y = 750
        for line in lines:
            c.drawString(100, y, line)
            y -= 50
        c.showPage()
    c.save()
    return buf.getvalue()


def _make_test_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, cell_text in enumerate(row):
                table.rows[i].cells[j].text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_pdf_extraction_gets_correct_text_and_page_count():
    pdf_bytes = _make_test_pdf([
        ["CONTRACT AGREEMENT", "Between Acme Corp and Beta Industries."],
        ["Termination clause applies after 12 months."],
    ])
    result = extract_pdf_text(pdf_bytes)
    assert result["page_count"] == 2
    assert "CONTRACT AGREEMENT" in result["text"]
    assert "Acme Corp" in result["text"]
    assert "Termination clause" in result["text"]
    assert result["word_count"] > 0


def test_docx_extraction_gets_paragraph_text():
    docx_bytes = _make_test_docx(["Employee Handbook", "Vacation policy: 20 days per year."])
    result = extract_docx_text(docx_bytes)
    assert "Employee Handbook" in result["text"]
    assert "Vacation policy" in result["text"]


def test_docx_extraction_includes_table_content():
    docx_bytes = _make_test_docx(
        ["Salary bands"], table_rows=[["Role", "Band"], ["Engineer", "Band 4"]]
    )
    result = extract_docx_text(docx_bytes)
    assert "Band 4" in result["text"], "table cell content must not be silently dropped"


def test_dispatch_routes_pdf_correctly():
    pdf_bytes = _make_test_pdf([["Some content here"]])
    result = extract_document_text(pdf_bytes, "report.pdf")
    assert result["document_type"] == "pdf"


def test_dispatch_routes_docx_correctly():
    docx_bytes = _make_test_docx(["Some content here"])
    result = extract_document_text(docx_bytes, "report.docx")
    assert result["document_type"] == "docx"


def test_corrupt_pdf_raises_extraction_error_not_a_crash():
    with pytest.raises(ExtractionError):
        extract_document_text(b"this is not a real pdf", "fake.pdf")


def test_corrupt_docx_raises_extraction_error_not_a_crash():
    with pytest.raises(ExtractionError):
        extract_document_text(b"garbage bytes", "fake.docx")


def test_unsupported_extension_raises_extraction_error():
    with pytest.raises(ExtractionError):
        extract_document_text(b"data", "notes.txt")
